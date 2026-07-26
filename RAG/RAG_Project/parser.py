import os
import re
from pypdf import PdfReader
import docx
from docx.oxml.ns import qn
import pdf_utils

def read_text_file(file_path):
    with open(file_path, 'r',encoding='utf-8') as f:
        text = f.read()
    return text

def read_md_file(file_path):
    """
    Reads a .md file as plain text, same mechanics as read_text_file.

    Kept as its own function (rather than just reusing read_text_file
    directly) because markdown needs DIFFERENT downstream handling --
    see the note in read_document() about why whitespace normalization
    is skipped for this format. Having a distinct function here also
    gives us a natural place to add markdown-specific logic later (e.g.
    stripping front-matter, or handling heading structure) without
    touching the plain-text path.
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()
    return text

def read_pdf_file(file_path, describe_images=True, image_backend="ocr"):
    """
    Reads a .pdf file into plain text, with cleanup aimed at real-world
    exports (Confluence, Word-to-PDF, etc.) rather than a clean single-
    column document:

      1. Extracts raw text per page.
      2. Strips repeating headers/footers (breadcrumbs, export dates,
         "Page N" markers) that pypdf would otherwise extract as if they
         were body content -- see pdf_utils.strip_boilerplate_lines().
      3. Skips pages that are a table of contents entirely -- dot-leader
         lines and duplicate heading text have no retrieval value and
         only add noise -- see pdf_utils.is_toc_page().
      4. If describe_images=True, extracts embedded images (diagrams,
         screenshots) and describes each one, inserting the description
         into the text where the image was. Without this, a diagram's
         actual content is invisible to the rest of the pipeline --
         extract_text() only ever returns text objects, never anything
         about an image sitting on the same page.

    Args:
        file_path: path to the PDF.
        describe_images: whether to caption embedded images at all.
            Set to False to skip image handling entirely (fastest, no
            dependencies).
        image_backend: which image-description method to use --
            "ocr"     -- fully offline, no model, just reads visible
                         text in the image (default; works out of the
                         box once tesseract-ocr is installed)
            "offline" -- fully offline, needs a local Ollama vision
                         model (better quality, understands diagram
                         structure, not just text) -- see
                         pdf_utils.describe_image_offline() for setup
            "claude"  -- needs internet + an API key, highest quality --
                         see pdf_utils.describe_image_with_claude()
    """
    reader = PdfReader(file_path)

    raw_page_texts = []
    for page in reader.pages:
        # extract_text() returns None (not "") for pages with no text
        # layer -- a blank page, a scanned/image-only page, etc.
        raw_page_texts.append(page.extract_text() or "")

    cleaned_page_texts = pdf_utils.strip_boilerplate_lines(raw_page_texts)

    # Build per-page image descriptions BEFORE assembling final text, so
    # we know which page index to attach each description to.
    image_descriptions_by_page = {}
    if describe_images:
        images = pdf_utils.extract_images(reader)
        for image in images:
            media_type = pdf_utils.guess_media_type(image["name"])
            try:
                description = pdf_utils.describe_image(
                    image["data"], media_type, backend=image_backend
                )
            except RuntimeError as e:
                # Missing package, missing local service, etc. -- don't
                # crash the whole parse over one image, just note it.
                description = f"[Image present but could not be described: {e}]"
            image_descriptions_by_page.setdefault(image["page_index"], []).append(description)

    final_pages = []
    for page_index, page_text in enumerate(cleaned_page_texts):
        if pdf_utils.is_toc_page(raw_page_texts[page_index]):
            continue  # skip the whole page -- TOC has no retrieval value

        page_parts = [page_text] if page_text.strip() else []

        for description in image_descriptions_by_page.get(page_index, []):
            page_parts.append(f"[Image: {description}]")

        if page_parts:
            final_pages.append("\n\n".join(page_parts))

    # joined with "\n\n" (not "") so the last word of one page and the
    # first word of the next don't get glued together with no space
    return "\n\n".join(final_pages)

def read_dox(file_path):
    """
    Reads headers, main body (paragraphs, tables, AND text boxes no
    matter how deeply nested), and footers.

    Why iter(qn("w:p")) instead of iterchildren()? iterchildren() only
    sees direct children of the body -- text inside a table cell or a
    text box lives nested several layers deep and gets silently missed.
    iter() searches the entire tree, so table cells and text boxes are
    caught the same way as ordinary paragraphs. The one thing to guard
    against: a paragraph that WRAPS a text box will itself contain
    another nested <w:p> -- if we counted both, the text would be
    duplicated, so wrapper paragraphs (any <w:p> containing a nested
    <w:p>) are skipped and only the innermost one is kept.
    """
    doc = docx.Document(file_path)
    parts = []
    # Header extraction
    for section in doc.sections:
        header_text = "\n".join(p.text for p in section.header.paragraphs if p.text.strip())
        if header_text:
            parts.append(header_text)

    # Main body extraction -- every paragraph anywhere in the tree,
    # in document order, regardless of nesting depth
    for p_elem in doc.element.body.iter(qn("w:p")):
        has_nested_paragraph = any(
            descendant is not p_elem for descendant in p_elem.iter(qn("w:p"))
        )
        if has_nested_paragraph:
            continue  # wrapper paragraph -- the real content is nested inside it

        run_texts = [t.text for t in p_elem.iter(qn("w:t")) if t.text]
        paragraph_text = "".join(run_texts).strip()
        if paragraph_text:
            parts.append(paragraph_text)

    # Footer extraction
    for section in doc.sections:
        footer_text = "\n".join(p.text for p in section.footer.paragraphs if p.text.strip())
        if footer_text:
            parts.append(footer_text)
    return "\n\n".join(parts)


def read_document(file_path, describe_images=True, image_backend="ocr"):
    extension = os.path.splitext(file_path)[1].lower()
    if extension == ".txt":
        text = read_text_file(file_path)
    elif extension == ".md":
        text = read_md_file(file_path)
    elif extension == ".pdf":
        text = read_pdf_file(file_path, describe_images=describe_images, image_backend=image_backend)
    elif extension == ".docx":
        text = read_dox(file_path)
    else:
        raise ValueError(f"Unsupported file type: {extension}")

    if extension == ".md":
        # Markdown is whitespace-sensitive: code blocks and nested lists
        # rely on exact indentation to mean what they mean. Collapsing
        # repeated spaces (the way we do for prose below) would flatten
        # a 4-space-indented code block and turn nested bullets into
        # siblings -- silently corrupting the document's structure. So
        # for markdown we skip that step and only trim outer whitespace.
        return text.strip()

    # clean up messy whitespace so chunking works on tidy text
    text = re.sub(r"[ \t]+"," ", text)   # Collapse repeated spaces/tabs
    text = re.sub(r"\n{3,}","\n\n", text)  # collapse 3+ blank lines to 1
    return text.strip()