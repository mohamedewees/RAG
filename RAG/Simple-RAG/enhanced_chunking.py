import os
import re
from pypdf import PdfReader
import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

def read_text_file(file_path):
    with open(file_path, 'r',encoding='utf-8') as f:
        text = f.read()
    return text

def read_pdf_file(file_path):
    reader = PdfReader(file_path)
    all_text = []
    for page in reader.pages:
        page_text = page.extract_text()
        all_text.append(page_text)
    return "".join(all_text)
    # return all_text

def read_dox(file_path):
    doc = docx.Document(file_path)
    parts = []
    # Header extraction
    for section in doc.sections:
        header_text = "\n".join(p.text for p in section.header.paragraphs if p.text.strip())
        if header_text:
            parts.append(header_text)
    # Main body extraction
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child,doc)
            if paragraph.text.strip():
                parts.append(paragraph.text)
        elif child.tag == qn("w:tbl"):
            table = Table(child,doc)
            for row in table.rows:
                cells_text = [cell.text.strip() for cell in row.cells]
                row_text = "|".join(t for t in cells_text if t)
                if row_text:
                    parts.append(row_text)
    
    # Footer extraction
    for section in doc.sections:
        footer_text = "\n".join(p.text for p in section.footer.paragraphs if p.text.strip())
        if footer_text:
            parts.append(footer_text)
    return "\n\n".join(parts)

def read_document(file_path):
    extension = os.path.splitext(file_path)[1].lower()
    if extension == ".txt":
        text = read_text_file(file_path)
    elif extension == ".pdf":
        text = read_pdf_file(file_path)
    elif extension == ".docx":
        text = read_dox(file_path)
    else:
        raise ValueError(f"Unsupported file type: {extension}")
    # clean up messy whitespace so chunking works on tidy text
    text = re.sub(r"[ \t]+"," ", text)   # Collapse repeated spaces/tabs
    text = re.sub(r"\n{3,}","\n\n", text)  # collapse 3+ blank lines to 1
    return text.strip()

SENTENCE_PATTERN = re.compile(r"(?<=[.!?])\s+(?=[A-Z])")
 
 
def split_into_sentences(paragraph):
    """Split a paragraph into sentences."""
    return [s.strip() for s in SENTENCE_PATTERN.split(paragraph) if s.strip()]
 
 
def _hard_split(text, max_len, overlap=50):
    """
    Safety-net splitter: force-splits a piece of text that's too long to
    fit in a chunk on its own, using a plain character sliding window
    (no sentence-awareness -- by the time we get here, sentence-awareness
    has already failed to produce something small enough).
    """
    pieces = []
    start = 0
    while start < len(text):
        end = min(start + max_len, len(text))
        piece = text[start:end].strip()
        if piece:
            pieces.append(piece)
        if end >= len(text):
            break
        start = end - overlap
    return pieces
 
 
def chunk_text(text, chunk_size=500, overlap_sentences=1):
    """
    Paragraph-aware, sentence-aware chunking with a guaranteed max size.
 
    Parameters
    ----------
    chunk_size : int
        Maximum characters per chunk. This is now a HARD ceiling --
        no chunk will exceed it, even for pathological input.
    overlap_sentences : int
        Number of pieces (sentences, or whole short paragraphs) to
        repeat at the start of the next chunk, for context continuity
        across the boundary.
    """
    if chunk_size <= 0:
        raise ValueError("chunk_size must be positive")
    if overlap_sentences < 0:
        raise ValueError("overlap_sentences cannot be negative")
 
    chunks = []
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
 
    current_pieces = []
    current_length = 0
 
    def flush(carry_overlap=True):
        """Emit the current chunk and optionally carry overlap forward."""
        nonlocal current_pieces, current_length
        if not current_pieces:
            return
        chunks.append(" ".join(current_pieces))
 
        if carry_overlap and overlap_sentences:
            current_pieces = current_pieces[-overlap_sentences:]
            current_length = sum(len(p) + 1 for p in current_pieces)
        else:
            current_pieces = []
            current_length = 0
 
    for paragraph in paragraphs:
        if len(paragraph) <= chunk_size:
            # Small enough to potentially sit as one unbroken piece.
            if current_length + len(paragraph) + 1 > chunk_size and current_pieces:
                flush()
            current_pieces.append(paragraph)
            current_length += len(paragraph) + 1
            continue
 
        # Paragraph is too big to add as one piece -- split into sentences.
        for sentence in split_into_sentences(paragraph):
            if len(sentence) > chunk_size:
                # SAFETY NET: even a single sentence is too long (e.g. a
                # bullet-separated line with no punctuation to split on).
                # Flush whatever we have, then hard-split just this piece.
                flush(carry_overlap=False)
                chunks.extend(_hard_split(sentence, chunk_size, overlap=50))
                continue
 
            if current_length + len(sentence) + 1 > chunk_size and current_pieces:
                flush()
            current_pieces.append(sentence)
            current_length += len(sentence) + 1
 
    flush(carry_overlap=False)  # emit whatever's left, nothing to carry into
 
    return chunks

def load_and_chunk(file_path , chunk_size = 200 , overlap = 50 , respect_sentences = True):
    text = read_document(file_path)
    # chunks = chunk_text(text , chunk_size=chunk_size , overlap=overlap , respect_sentences=respect_sentences)
    chunks = chunk_text(text)
    return chunks

if __name__ == "__main__":
    file_path = r"/home/mohamed-ewees/Downloads/My CV-20260716T121707Z-1-001/My CV/DevOps/AI generated/Mohammed_Ewees.docx"
    chunks = load_and_chunk(file_path, chunk_size=200 , overlap=50)
    print(f"Loaded and chunked {file_path} into {len(chunks)} chunks:\n")

    for i, chunk in enumerate(chunks):
        print(f"--- Chunk {i} ({len(chunk)} chars) ---")
        print(chunk)
        print()
